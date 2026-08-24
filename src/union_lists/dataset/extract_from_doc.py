from multiprocessing import Value
from collections import Counter
from datetime import datetime
import glob
import json
import logging
import os
import re
import shutil
import xml.etree.ElementTree as ET
import xml.dom.minidom

from docx import Document, table
import numpy as np
import pandas as pd
from tqdm import tqdm
import win32com.client as win32

from union_lists.config import INTERIM_DATA_DIR, LOGS_DIR

logger = logging.getLogger(__name__)
logging.basicConfig(filename=f"{LOGS_DIR}/{datetime.now().strftime("%Y%m%d_%H%M")}_extract.log", level=logging.INFO)


def find_refs(text) -> list[str]:
    """Extract X/ or W/ references from a string

    Args:
        text (str): Text from a union list document that may contain X/ or W/ references

    Returns:
        list[str]: A list of extracted refs
    """
    ref_re = re.compile(r"(?<![NS])[XW]/.+?1[89][0-9]{2,2}(?!/)", re.DOTALL)
    refs = [ref.strip().strip("\n") for ref in ref_re.findall(text)]
    return refs


def extract_label_lines(paragraphs):
    """Label lines as references or notes based on use of italic
    Any lines with the MS Word property of italic are notes
    Any in standard format are references

    Args:
        paragraphs (docx.text.paragraph.Paragraph): _description_

    Returns:
        str: line break escaped text lines with trailing `/// <ref/note>`
    """
    ref_start_re = re.compile(r"(?<![NS])[XW]/")

    lines = []
    skip_para = False
    for i, para in enumerate(paragraphs):
        
        if skip_para:
            skip_para = False
            continue
            
        non_space_runs = [run for run in para.runs if run.text not in [" ", "  ", "\n"]]
        itals = [run.italic for run in non_space_runs]
        if set(itals) == {False, None}:  # There are a few cells with italic = None and non italic = False (for some godforsaken reason)
            for run in non_space_runs:
                if run.italic == False:
                    run.italic = None
                elif run.italic is None:
                    run.italic = True
        elif set(itals) == {False, None, True}:
            for run in non_space_runs:
                if run.italic == False:
                    run.italic = None

        ital_count = sum([run.italic for run in non_space_runs if run.italic])
        
        # breakpoint()
        if not non_space_runs:
            continue

        # TODO finish implementing finding references where there are notes not in italic (for some reason)
        refs = find_refs(para.text)
        ref_start = ref_start_re.search(para.text)
        # breakpoint()
        if not non_space_runs[0].italic and ref_start and not refs and i < len(paragraphs) - 1:
            # must not start with italic
            # must contain the start of a reference
            # must not contain a full reference
            # must not be the final line
            logging.info("ref line split over multiple lines")
            para.text = para.text.strip() + " " + paragraphs[i+1].text.strip()
            skip_para = True
            refs = find_refs(para.text)

        elif ital_count == 0 and len(refs) == 1:
            lines.append(f"{para.text} /// ref")

        elif ital_count == 0 and len(refs) > 1:
            logging.info(f"ref line with > 1 ref found: {len(refs)} refs, {para.text}")
            reformed_ref_lines = [ref + " /// ref" for ref in refs]
            lines.extend(reformed_ref_lines)
            
        elif len(refs) == 0 and ital_count == 0:
            logging.info(f"no ital line with 0 refs found: {para.text}")
            lines.append(f"{para.text} /// note")

        elif ital_count and (ital_count == len(non_space_runs)):
            lines.append(f"{para.text} /// note")

        else:
            line_tracker = {None: [], True: []}
            ital_tracker = non_space_runs[0].italic
            end = {None: " /// ref", True: " /// note"}
            for i, run in enumerate(non_space_runs):
                final_line = i == len(non_space_runs) - 1
                # breakpoint()

                if run.italic != ital_tracker and ital_tracker is None:
                    # non-ital > ital, extract refs
                    full_line = " ".join(line_tracker[ital_tracker])
                    refs = find_refs(full_line)
                    reformed_ref_lines = [ref + " /// ref" for ref in refs]

                    lines.extend(reformed_ref_lines)

                    line_tracker[ital_tracker] = []
                    line_tracker[run.italic].append(run.text)
                    ital_tracker = run.italic

                    if final_line:
                        ital_line = run.text.strip("\n").strip() + " /// note"
                        lines.append(ital_line)

                elif run.italic != ital_tracker and ital_tracker:
                    # ital > non-ital
                    # breakpoint()
                    new_line_stripped = [l.strip("\n") for l in line_tracker[ital_tracker]]
                    ital_line = " ".join(new_line_stripped) + " /// note"
                    lines.append(ital_line)

                    line_tracker[ital_tracker] = []
                    line_tracker[run.italic].append(run.text)
                    ital_tracker = run.italic

                    if final_line:
                        ref_line = run.text.strip("\n").strip() + " /// ref"
                        lines.append(ref_line)

                elif final_line and ital_tracker is None:
                    # final line, non-ital so extract refs
                    # breakpoint()
                    line_tracker[run.italic].append(run.text)
                    combined_ref_lines = " ".join(line_tracker[ital_tracker]) + end[ital_tracker]

                    refs = find_refs(combined_ref_lines)
                    reformed_ref_lines = [ref + " /// ref" for ref in refs]
                    lines.extend(reformed_ref_lines)

                elif final_line and ital_tracker:
                    # final line, ital so append combined
                    line_tracker[run.italic].append(run.text)
                    new_line_stripped = [l.strip("\n") for l in line_tracker[ital_tracker]]
                    line = " ".join(new_line_stripped) + " /// note"
                    lines.append(line)

                else:
                    # continuation
                    line_tracker[run.italic].append(run.text)

    lines = [l for l in lines if l != " /// ref"]

    return "\n".join(lines)


def extract_headers(row):
    headers = {
        "Post-1905": [[], None],
        "1886-1905": [[], None],
        "Pre-1886": [[], None]
    }

    col_names = []

    post_1905_suffix = 1
    mid_year_suffix = 1
    pre_1886_suffix = 1

    for i, cell in zip(row.index, row):
        if "Post-1905" in cell:
            headers["Post-1905"][0].append(i)  # ty:ignore[unresolved-attribute]
            headers["Post-1905"][1] = cell
            col_names.append(f"Post-1905_{post_1905_suffix}")
            post_1905_suffix += 1
        elif "1886-" in cell:
            headers["1886-1905"][0].append(i)  # ty:ignore[unresolved-attribute]
            headers["1886-1905"][1] = cell
            col_names.append(f"1886-1905_{mid_year_suffix}")
            mid_year_suffix += 1
        elif "Pre-1886" in cell:
            headers["Pre-1886"][0].append(i)  # ty:ignore[unresolved-attribute]
            headers["Pre-1886"][1] = cell
            col_names.append(f"Pre-1886_{pre_1886_suffix}")
            pre_1886_suffix += 1
        else:
            raise ValueError(f"No date information in cell: {cell}")

    return headers, col_names
        

def clean_map_df(table_dict: dict[int, list[table._Cell]]) -> tuple[pd.DataFrame, dict[str, str]]:
    cell_df = pd.DataFrame(table_dict)
    df = cell_df.apply(lambda x: x.apply(lambda y: y.text))
    
    dup_cols = []
    for col in df.columns[:-1]:
        # Skip the first couple of values
        # Post-1905 is sometimes duplicated apart from the first row containing the heading block number
        if np.array_equal(df[col].values[2:], df[col+1].values[2:]):  
            dup_cols.append(col + 1)

    df = df.drop(columns=dup_cols)
    cell_df = cell_df.drop(columns=dup_cols)

    if len(df) == 6 and df.apply(lambda x: x.str.contains("are known to have")).any().any():
        no_known_text = np.unique(df.apply(lambda x: x[x.str.contains("are known to have")]).values)[0]
        df = pd.DataFrame()
        table_metadata = {"no_known_maps": no_known_text}

    elif df.shape[1] == 2:
        df = df.dropna(how="all").reset_index(drop=True)
        df.columns = ["Post-1905_1", "metadata"]
        table_metadata = {}
        df = df.apply(lambda x: x.str.strip().str.rstrip("\n"), axis=1)

    elif df.shape[1] > 6 or df.shape[1] == 3:
        raise ValueError(f"df has unexpected shape {df.shape}")

    elif df.shape[1] <= 6 and df.shape[1] > 3:
        long_lines_idx = df.apply(lambda x: x.transform(len)).sum(axis=1).sort_values(ascending=False).iloc[:2].index
        header_row, footer_row = long_lines_idx.min(), long_lines_idx.max()
        headers, col_names = extract_headers(df.loc[header_row])
        df.columns = col_names
            
        post_1905_header, mid_header, pre_1886_header = headers["Post-1905"][1], headers["1886-1905"][1], headers["Pre-1886"][1]
        post_1905_footer, mid_footer, pre_1886_footer = df.loc[footer_row,["Post-1905_1", '1886-1905_1', 'Pre-1886_1']]
        table_metadata = {
            "Post-1905_header": post_1905_header, "1886-1905_header": mid_header, "Pre-1886_header": pre_1886_header,
            "Post-1905_footer": post_1905_footer, "1886-1905_footer": mid_footer, "Pre-1886_footer": pre_1886_footer
        }
        
        col_2s = [c for c in col_names if "_2" in c]
        missing_cols = f"{"1886-1905_2" * ("1886-1905_2" in col_2s)} {"Pre-1886_2" * ("Pre-1886_2" in col_2s)}"
        
        if df.shape[1] == 4:
            if "1886-1905_2" in col_2s or "Pre-1886_2" in col_2s:
                raise ValueError(f"Unexpected column in the bagging area: Missing columns {missing_cols}")
            logging.info(f"{missing_cols} missing from column list")
        elif df.shape[1] == 5:
            if "1886-1905_2" in col_2s and "Pre-1886_2" in col_2s:
                raise ValueError(f"Unexpected column in the bagging area: Missing columns {missing_cols}")
            logging.info(f"{missing_cols} missing from column list")
        elif df.shape[1] == 6:
            if "1886-1905_2" not in col_2s and "Pre-1886_2" not in col_2s:
                raise ValueError(f"Unexpected column in the bagging area: Missing columns {missing_cols}")
            logging.info(f"nothing missing from column list")

        df = df.reindex(columns=['Post-1905_1', 'Post-1905_2', '1886-1905_1', '1886-1905_2', 'Pre-1886_1', 'Pre-1886_2'])
        col_map = {"Post-1905_2": 1, "1886-1905_2": 3, "Pre-1886_2": 4 + ("1886-1905_2" in col_2s)}
        for c in col_2s:
            df[c] = cell_df.iloc[df.index, col_map[c]].apply(lambda cell: extract_label_lines(cell.paragraphs))

        df = df.where(lambda x: x != '', pd.NA).dropna(how="all", axis=0)
        iloc_header_row, iloc_footer_row = df.index.to_list().index(header_row), df.index.to_list().index(footer_row)
        
        if header_row > 9 or len(df) - footer_row > 10:
            raise ValueError(f"Header or footer in unlikely location header: {header_row}, footer: {footer_row}")
        
        df = df.iloc[iloc_header_row + 1: iloc_footer_row].copy()
        df = df.reset_index(drop=True)

    df = df.apply(lambda x: x.str.strip().str.rstrip("\n"), axis=1)
    return df, table_metadata


def postcorrect_df(df, file_id):
    if file_id == "38A":
        # ref missing date, add 9999
        assert df.loc[172, "Post-1905_2"] == 'X/9053/38K/13  1908 /// ref\nBazar Valley area only of Khyber Agency /// note\nX/9053/38K/13 1921 /// ref\nX/9053/38K/13  1921/1928 /// ref\nX/9053/38K/13  1934/1935 /// ref'
        df.loc[172, "Post-1905_2"] = 'X/9053/38K/13  1908 /// ref\nBazar Valley area only of Khyber Agency /// note\nX/9053/38K/13  1921 /// ref\nX/9053/38K/13  1921/1928 /// ref\nX/9053/38K/13  1934/1935 /// ref\nX/9053/38K/13+J/16 9999 /// ref\npart extended north for Bazar Valley area of Khyber Agency /// note'
    
    if file_id == "44A":
        # space separating ref and date has been replaced with a '/'
        assert df.loc[75, "1886-1905_2"] == 'X/9373/195  1886 /// ref\nMontgomery and Shekhupura only /// note\nX/9373/195/1904 /// ref\nMontgomery, Lahore, and part only of Shekhupura /// note'
        df.loc[75, "1886-1905_2"] = 'X/9373/195  1886 /// ref\nMontgomery and Shekhupura only /// note\nX/9373/195 1904 /// ref\nMontgomery, Lahore, and part only of Shekhupura /// note'

    if file_id == "62A":
        # runs split words
        assert df.loc[48, "Post-1905_2"] == 'X/9053/62D /1+ D/ 5  1924 /// ref\nUnited Provinces  only /// note'
        df.loc[48, "Post-1905_2"] = 'X/9053/62D/1+D/5  1924 /// ref\nUnited Provinces  only /// note'

    if file_id == "65A":
        # '/' replaced with a '.'
        assert df.loc[175, "Post-1905_2"] == 'X/9053/65L.5  1941 /// ref'
        df.loc[175, "Post-1905_2"] = 'X/9053/65L/5  1941 /// ref'

    if file_id == "72A":
        # first ref missing - unknown reason
        assert df.loc[93, "Post-1905_2"] == 'Bihar  only /// note'
        df.loc[93, "Post-1905_2"] = 'X/9053/72F/14  1939 /// ref\nBihar only /// note'

    if file_id == "73C":
        # errant space in reference
        assert df.loc[11, "Post-1905_2"] == 'X/9051/73L+ P  1912 /// ref\npart extended east for coastal area /// note\nX/9051/73L+P+74 I  1934 /// ref\nparts extended east and south for coastal area /// note\nX/13104/73L+P+74 I  1934 /// ref\nparts extended east and south for coastal area /// note'
        df.loc[11, "Post-1905_2"] = 'X/9051/73L+P 1912 /// ref\npart extended east for coastal area /// note\nX/9051/73L+P+74I 1934 /// ref\nparts extended east and south for coastal area /// note\nX/13104/73L+P+74I 1934 /// ref\nparts extended east and south for coastal area /// note'
    return df


if __name__ == "__main__":
    ns_raw = 'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex" xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex" xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex" xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex" xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex" xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex" xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex" xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du" xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" xmlns:w16sdtfl="http://schemas.microsoft.com/office/word/2024/wordml/sdtformatlock" xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'.split()
    ms_ns = {ns.split("=")[0][6:]: ns.split("=")[1].strip('"') for ns in ns_raw}
    [ET.register_namespace(prefix, uri) for prefix, uri in ms_ns.items()];

    SCALE = "Quarter Inch"

    docx_files = glob.glob(f"{INTERIM_DATA_DIR}/{SCALE}/*.docx")
    docx_files = [x for x in docx_files if "\\~" not in x and "(2)" not in x]
    docx_files = [x for x in docx_files if "_mod" not in x]
    scale_docs = {"Quarter Inch": 45, "Half Inch": 38, "One Inch": 44}
    assert len(docx_files) == scale_docs[SCALE]

    print(f"{SCALE}: Extracting text and cleaning tables")
    empty_docs = []
    with tqdm(docx_files) as t:
        for f in t:
            doc_table = Document(f).tables[0]
            table_dict = {i: doc_table.column_cells(i) for i, _ in enumerate(doc_table.columns)}
            file_id = os.path.basename(f).split(".")[0]
            t.set_description(file_id)
            df, metadata = clean_map_df(table_dict)
            df = postcorrect_df(df=df, file_id=file_id)

            if not df.empty:
                df.to_csv(f"{INTERIM_DATA_DIR}/{SCALE}/{file_id}.csv", encoding="utf-8-sig", index=False)
            else:
                empty_docs.append(file_id)
            
            # Take care as this formulation means number of output JSONs can be larger than num CSVs
            with open(f"{INTERIM_DATA_DIR}/{SCALE}/{file_id}.json", "w",encoding="utf8") as f:
                json.dump(metadata, f)
    
    with open(f"{INTERIM_DATA_DIR}/{SCALE}/empty_doc.log", "w", encoding="utf8") as f:
        [f.write(f"{file_id} is empty\n") for file_id in empty_docs]