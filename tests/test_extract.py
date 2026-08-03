import glob
import xml.etree.ElementTree as ET

from docx import Document, table
import pandas as pd
import pytest
from pytest import raises

from union_lists.dataset import extract_from_doc as extract

ns_raw = 'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex" xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex" xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex" xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex" xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex" xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex" xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex" xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du" xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" xmlns:w16sdtfl="http://schemas.microsoft.com/office/word/2024/wordml/sdtformatlock" xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'.split()
ms_ns = {ns.split("=")[0][6:]: ns.split("=")[1].strip('"') for ns in ns_raw}
[ET.register_namespace(prefix, uri) for prefix, uri in ms_ns.items()];


def test_find_refs():
    partial_ref = "X/9052/37P/SE+P/NE"
    assert extract.find_refs(partial_ref) == []

    one_ref = "X/9051/58B  1922"
    assert extract.find_refs(one_ref) == ["X/9051/58B  1922"]

    two_refs = "X/9051/58B  1922/1937 X/13104/58B  1922/1944"
    assert extract.find_refs(two_refs) == ["X/9051/58B  1922/1937", "X/13104/58B  1922/1944"]

    x_w_refs = "X/9051/62B  1881/1923\nX/9051/62B  1931\nW/LPS/21/N4/62B  1931\nX/13104/62B  1943/1953"
    assert extract.find_refs(x_w_refs) == ["X/9051/62B  1881/1923", "X/9051/62B  1931", "W/LPS/21/N4/62B  1931", "X/13104/62B  1943/1953"]


def test_extract_label_lines():

    invert_ital_ref_table = Document("tests/62A.docx").tables[0]
    invert_ital_ref_dict = {i: invert_ital_ref_table.column_cells(i) for i, _ in enumerate(invert_ital_ref_table.columns)}
    invert_ital_ref_df = pd.DataFrame(invert_ital_ref_dict)

    invert_ital_ref_result = extract.extract_label_lines(invert_ital_ref_df.iloc[54, 1].paragraphs)
    invert_ital_ref_expected = 'X/9053/62D /1+ D/ 5  1924 /// ref\nUnited Provinces  only /// note'

    assert invert_ital_ref_result == invert_ital_ref_expected

    false_trailing_ref_table = Document("tests/53A.docx").tables[0]
    false_trailing_ref_dict = {i: false_trailing_ref_table.column_cells(i) for i, _ in enumerate(false_trailing_ref_table.columns)}
    false_trailing_ref_df = pd.DataFrame(false_trailing_ref_dict)

    false_trailing_ref_result = extract.extract_label_lines(false_trailing_ref_df.iloc[153, 7].paragraphs)
    false_trailing_ref_expected = 'see list for X/9322/1-4 /// note\nX/1493/1/1  1867 /// ref\nGarhwal only /// note\nX/1493/3/1  1876 /// ref\nGarhwal only; skeleton sheet /// note'

    assert false_trailing_ref_result == false_trailing_ref_expected

    trailing_ref_table = Document("tests/39B.docx").tables[0]
    trailing_ref_dict = {i: trailing_ref_table.column_cells(i) for i, _ in enumerate(trailing_ref_table.columns)}
    trailing_ref_df = pd.DataFrame(trailing_ref_dict)

    trailing_ref_result = extract.extract_label_lines(trailing_ref_df.iloc[25, 4].paragraphs)
    trailing_ref_expected = 'X/9935/2/454+458  1885 /// ref\n29NE/3+29NW/4: Dera Ismail Khan, and part only of Loralai and Zhob, 69°40′E-70°40′E /// note\nX/9935/2/454  1901 /// ref'

    assert trailing_ref_result == trailing_ref_expected

    trailing_ital_table = Document("tests/64B.docx").tables[0]
    trailing_ital_dict = {i: trailing_ital_table.column_cells(i) for i, _ in enumerate(trailing_ital_table.columns)}
    trailing_ital_df = pd.DataFrame(trailing_ital_dict)

    trailing_ital_result = extract.extract_label_lines(trailing_ital_df.iloc[61, 1].paragraphs)
    trailing_ital_expected = 'X/9052/64N/SE+O/NE  1916 /// ref\ncurtailed north at 22°20’N, part extended north of 22°20’N for Gangpur only; extended south to 21°50’N for Gangpur, part extruded further south for Gangpur only see also 73B/SW /// note'

    assert trailing_ital_result == trailing_ital_expected

    see_doc_table = Document("tests/41C.docx").tables[0]
    see_table_dict = {i: see_doc_table.column_cells(i) for i, _ in enumerate(see_doc_table.columns)}
    see_cell_df = pd.DataFrame(see_table_dict)

    see_result = extract.extract_label_lines(see_cell_df.iloc[14, 1].paragraphs)
    see_expected = 'see 41K+L /// note'

    assert see_result == see_expected
    
    no_content_doc_table = Document("tests/42C.docx").tables[0]
    no_content_table_dict = {i: no_content_doc_table.column_cells(i) for i, _ in enumerate(no_content_doc_table.columns)}
    no_content_cell_df = pd.DataFrame(no_content_table_dict)

    no_content_result = extract.extract_label_lines(no_content_cell_df.iloc[0, 3].paragraphs)
    no_content_expected = ''

    assert no_content_result == no_content_expected

    qi_doc_table = Document("tests/42C.docx").tables[0]
    qi_table_dict = {i: qi_doc_table.column_cells(i) for i, _ in enumerate(qi_doc_table.columns)}
    qi_cell_df = pd.DataFrame(qi_table_dict)

    qi_result = extract.extract_label_lines(qi_cell_df.iloc[8, 3].paragraphs)
    qi_expected = 'X/9936/1/2NW  1896 /// ref\nAfghanistan, and part only of Russian Turkestan /// note\nX/9936/1/2NW  1907 /// ref'

    assert qi_result == qi_expected

    two_ref_line_doc_table = Document("tests/58C.docx").tables[0]
    two_ref_line_table_dict = {i: two_ref_line_doc_table.column_cells(i) for i, _ in enumerate(two_ref_line_doc_table.columns)}
    two_ref_line_cell_df = pd.DataFrame(two_ref_line_table_dict)

    two_ref_line_result = extract.extract_label_lines(two_ref_line_cell_df.iloc[7, 1].paragraphs)
    two_ref_line_expected = 'X/9051/58B  1922 /// ref\nX/9051/58B  1922/1937 /// ref\nX/13104/58B  1922/1944 /// ref'
    
    assert two_ref_line_result == two_ref_line_expected
    
    long_note_doc_table = Document("tests/63C.docx").tables[0]
    long_note_table_dict = {i: long_note_doc_table.column_cells(i) for i, _ in enumerate(long_note_doc_table.columns)}
    long_note_cell_df = pd.DataFrame(long_note_table_dict)
    # breakpoint()
    long_note_result = extract.extract_label_lines(long_note_cell_df.iloc[20, 1].paragraphs)
    long_note_expected = "X/9051/63O 1921 /// ref\nX/9051/63O 1933 /// ref\npart ext en ded west for Benares City only /// note\nX/13104/63O 1948 /// ref\npart ext en ded west for Benares City only /// note"

    assert long_note_result == long_note_expected

    plus_doc_table = Document("tests/73C.docx").tables[0]
    plus_table_dict = {i: plus_doc_table.column_cells(i) for i, _ in enumerate(plus_doc_table.columns)}
    plus_cell_df = pd.DataFrame(plus_table_dict)
    
    plus_result = extract.extract_label_lines(plus_cell_df.iloc[17, 1].paragraphs)
    plus_expected = 'X/9051/73L+P  1912 /// ref\npart extended east for coastal area /// note\nX/9051/73L+P+74I  1934 /// ref\nparts extended east and south for coastal area  /// note\nX/13104/73L+P+74I  1934/1942 /// ref\nparts extended east and south for coastal area /// note'
    assert plus_result == plus_expected

    multi_final_ref_doc_table = Document("tests/78C.docx").tables[0]
    multi_final_ref_table_dict = {i: multi_final_ref_doc_table.column_cells(i) for i, _ in enumerate(multi_final_ref_doc_table.columns)}
    multi_final_ref_df = pd.DataFrame(multi_final_ref_table_dict)

    multi_final_ref_result = extract.extract_label_lines(multi_final_ref_df.iloc[6, 1].paragraphs)
    multi_final_ref_expected = 'X/9051/78A+77D  1921 /// ref\nextended north for  Sikkim /// note\nX/9051/78A+77D  1923 /// ref\nextended north for  Sikkim /// note\nX/9051/78A  1937 /// ref\nW/LPS/21/N4/78A  1937 /// ref'

    assert multi_final_ref_result == multi_final_ref_expected


def test_postcorrect_df():
    file_id = "62A"
    postcorr_doc_table = Document(f"tests/{file_id}.docx").tables[0]
    postcorr_table_dict = {i: postcorr_doc_table.column_cells(i) for i, _ in enumerate(postcorr_doc_table.columns)}
    
    postcorr_df, metadata = extract.clean_map_df(table_dict=postcorr_table_dict)
    postcorr_df = extract.postcorrect_df(df=postcorr_df, file_id=file_id)

    postcorr_expected = 'X/9053/62D/1+D/5  1924 /// ref\nUnited Provinces  only /// note'
    assert postcorr_df.loc[48, "Post-1905_2"] == postcorr_expected
